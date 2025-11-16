import os
import sys
import base64
import mimetypes
import zipfile
from pathlib import Path
import tkinter as tk
from tkinter import ttk, filedialog, simpledialog, messagebox, scrolledtext
import threading
import json

from openai import OpenAI
from docx import Document
import fitz 
import sv_ttk 
from PIL import Image, ImageTk

# ----------------------------------------------------------------------
# PATHS
# ----------------------------------------------------------------------
def get_base_dir():
    if getattr(sys, 'frozen', False):
        return sys._MEIPASS
    else:
        return os.path.dirname(os.path.abspath(__file__))

BASE_DIR = get_base_dir()
HELP_FILES = [
    os.path.join(BASE_DIR, "HELP.txt"),
]
ACK_FILES = [
    os.path.join(BASE_DIR, "ACKNOWLEDGEMENT.txt"),
]
LICENSE_FILES = [
    os.path.join(BASE_DIR, "LICENSE.txt"),
]
USER_PHOTO_FILES = [
    os.path.join(BASE_DIR, "user_photo.png"),
]
USER_ICON_FILE = [
    os.path.join(BASE_DIR, "user_photo.ico")
]

# ----------------------------------------------------------------------
# FILE LOADERS
# ----------------------------------------------------------------------

def load_first_existing(paths, fallback_text):
    for p in paths:
        try:
            if os.path.exists(p):
                with open(p, "r", encoding="utf-8") as f:
                    return f.read()
        except Exception as e:
            print(f"Error reading file {p}: {e}")
            continue 
    return fallback_text

HELP_TEXT = load_first_existing(
    HELP_FILES,
    "Help.txt not found. Please place 'HELP.txt' in the application folder.\n\n"
)

ACK_TEXT = load_first_existing(
    ACK_FILES,
    "Acknowledgment file (ACKNOWLEDGEMENT.txt) not found.\n"
)

LICENSE_TEXT = load_first_existing(
    LICENSE_FILES,
    "LICENSE.txt not found.",
)

# ----------------------------------------------------------------------
# STYLED UI HELPER
# ----------------------------------------------------------------------

def add_photo_header(window, title, subtitle=None):
    header = ttk.Frame(window, padding="20 10 20 20")
    header.pack(fill=tk.X, side=tk.TOP)

    root = window.master
    user_photo = getattr(root, "user_photo_small", None)

    header.columnconfigure(1, weight=1)

    if user_photo is not None:
        img_label = ttk.Label(header, image=user_photo)
        img_label.image = user_photo
        img_label.grid(row=0, column=0, rowspan=2, sticky=tk.NW, padx=(0, 15))
    
    title_label = ttk.Label(header, text=title, font=("Segoe UI", 14, "bold"))
    title_label.grid(row=0, column=1, sticky=tk.W)

    if subtitle:
        subtitle_label = ttk.Label(
            header, text=subtitle, font=("Segoe UI", 10)
        )
        subtitle_label.grid(row=1, column=1, sticky=tk.W, pady=(2, 0))

# ----------------------------------------------------------------------
# POP-UP WINDOWS (From Clinical Anonymizer)
# ----------------------------------------------------------------------

class HelpWindow:
    def __init__(self, parent):
        self.parent = parent

    def show(self):
        w = tk.Toplevel(self.parent)
        w.title("Usage Help – AI Clinical Interpreter")
        w.geometry("700x520")
        w.transient(self.parent)
        w.grab_set()
        if hasattr(self.parent, "user_photo_small"):
            w.iconphoto(False, self.parent.user_photo_small)

        add_photo_header(
            w,
            "Usage Help",
            "How to use the AI Clinical Interpreter.",
        )

        main = ttk.Frame(w, padding="0 20 20 20")
        main.pack(fill=tk.BOTH, expand=True)
        main.rowconfigure(0, weight=1)
        main.columnconfigure(0, weight=1)

        txt = scrolledtext.ScrolledText(main, wrap=tk.WORD, width=80, height=20, font=("Segoe UI", 10))
        txt.grid(row=0, column=0, sticky="nsew")
        txt.insert(tk.END, HELP_TEXT)
        txt.config(state=tk.DISABLED)

        btn_frame = ttk.Frame(main)
        btn_frame.grid(row=1, column=0, sticky=tk.E, pady=(15, 0))
        ttk.Button(btn_frame, text="Close", command=w.destroy).pack()

class AcknowledgmentWindow:
    def __init__(self, parent):
        self.parent = parent

    def show(self):
        w = tk.Toplevel(self.parent)
        w.title("Acknowledgments – AI Clinical Interpreter")
        w.geometry("700x520")
        w.transient(self.parent)
        w.grab_set()
        if hasattr(self.parent, "user_photo_small"):
            w.iconphoto(False, self.parent.user_photo_small)

        add_photo_header(
            w,
            "Acknowledgments / Danksagung",
            "Support for the development of this tool.",
        )

        main = ttk.Frame(w, padding="0 20 20 20")
        main.pack(fill=tk.BOTH, expand=True)
        main.rowconfigure(0, weight=1)
        main.columnconfigure(0, weight=1)

        txt = scrolledtext.ScrolledText(main, wrap=tk.WORD, width=80, height=20, font=("Segoe UI", 10))
        txt.grid(row=0, column=0, sticky="nsew")
        txt.insert(tk.END, ACK_TEXT)
        txt.config(state=tk.DISABLED)

        btn_frame = ttk.Frame(main)
        btn_frame.grid(row=1, column=0, sticky=tk.E, pady=(15, 0))
        ttk.Button(btn_frame, text="Close", command=w.destroy).pack()

class LicenseWindow:
    def __init__(self, parent):
        self.parent = parent

    def show(self):
        w = tk.Toplevel(self.parent)
        w.title("View License – AI Clinical Interpreter")
        w.geometry("700x520")
        w.transient(self.parent)
        w.grab_set()
        if hasattr(self.parent, "user_photo_small"):
            w.iconphoto(False, self.parent.user_photo_small)

        add_photo_header(
            w,
            "License",
            "GPL license and conditions of use.",
        )

        main = ttk.Frame(w, padding="0 20 20 20")
        main.pack(fill=tk.BOTH, expand=True)
        main.rowconfigure(0, weight=1)
        main.columnconfigure(0, weight=1)

        txt = scrolledtext.ScrolledText(main, wrap=tk.WORD, width=80, height=20, font=("Segoe UI", 10))
        txt.grid(row=0, column=0, sticky="nsew")
        txt.insert(tk.END, LICENSE_TEXT)
        txt.config(state=tk.DISABLED)

        btn_frame = ttk.Frame(main)
        btn_frame.grid(row=1, column=0, sticky=tk.E, pady=(15, 0))
        ttk.Button(btn_frame, text="Close", command=w.destroy).pack()

# ----------------------------------------------------------------------
# PROMPTS
# ----------------------------------------------------------------------

CLINICAL_PROMPT = """
ROLE:
You are an expert-level Clinical Motion Scientist with deep specialization in orthopedics, movement analysis, and rehabilitation. You are assisting a busy clinician (e.g., surgeon, physiatrist, physiotherapist) who needs a short, decision-oriented report.

CONTEXT:
Your task is to integrate all provided anonymized files for ONE patient. Files may include gait and pressure data, kinematics, kinetics, spatiotemporal parameters, EMG, range of motion, strength tests, patient-reported outcomes, or imaging summaries. The user has provided a description for each file.

GENERAL RULES:
- Focus on PATTERNS and CLINICALLY RELEVANT abnormalities, not on listing every parameter.
- Avoid repetition between sections. If something is already mentioned (e.g., in KEY ABNORMALITIES), refer back to it instead of rewriting the details.
- Write in clear, professional clinical English.
- NEVER include or infer any Patient Health Information (PHI). Refer to the subject only as “the patient”.
- TOTAL LENGTH: aim for about 400–600 words maximum. Shorter is preferred if the data are simple.

REFERENCE NORMS:
- Use any lab-specific ranges present in the files first.
- If none are present, use widely accepted healthy adult norms (e.g., comfortable gait speed ≈ 1.2–1.4 m/s) OR norms explicitly stated by the user.
- Only mention norms when they help explain an abnormality. Avoid overloading the text with numbers.
- Always state the norm you are using when you call something “increased”, “reduced”, or “abnormal”.

OUTPUT FORMAT
(Use these EXACT headings and order. Do NOT add extra headings. Keep each section concise.)

1) EXECUTIVE SUMMARY
Provide 3–5 bullet points that a busy clinician can read in <30 seconds.
- Each bullet = 1 sentence.
- Summarize the main pattern, severity (mild/moderate/severe), and overall functional impact.
- Do NOT include long lists of numbers here.

2) OBJECTIVE DATA SNAPSHOT
Summarize the most important objective findings from all files, grouped by source.
- Max 8 bullets total.
- Use the user’s description as a label for each file.
- Format:
  - [Description label]: <very short phrase about pattern> – key values only when needed, e.g.
    “Gait speed 0.8 m/s (norm ~1.2–1.4), double support 41% (high), stance longer on left (73% vs 68%).”
- Highlight only abnormalities or key contextual values (speed, cadence, symmetry). Do NOT list every parameter.

3) KEY ABNORMALITIES (Numbered List)
List the top 3–5 clinically most relevant findings, each on ONE line with this structure:
- 1. Finding -> Likely mechanism / cause (if reasonably inferable) -> Functional impact
Examples:
- 1. Prolonged left stance and single support -> likely right-side discomfort/weakness or instability -> increased load on left limb, mild gait asymmetry.
- 2. High double support time at slow speed -> cautious, stability-seeking strategy -> reduced dynamic balance demands but slower, inefficient gait.
If the mechanism is uncertain, write “possible mechanisms: …” instead of inventing a precise cause.

4) CLINICAL CONCLUSION & DECISION SUPPORT
Provide a single short paragraph (3–6 sentences).
- Integrate the KEY ABNORMALITIES into a coherent clinical picture (overall pattern, likely main driver(s), and severity).
- Comment briefly on:
  - Functional impact (e.g., endurance, community ambulation, stairs, uneven ground).
  - Joint loading / overuse risk (which side, which joints) if relevant.
  - Fall risk / balance concerns if suggested by the data.
- Do NOT repeat all the numbers; instead, refer to them conceptually (“temporal asymmetry with prolonged left stance”, “high double support”, etc.).

5) RECOMMENDED NEXT STEPS (Checklist)
Give a very concise, prioritized checklist of next steps (max 6 items total).
Use this structure and keep each item to ONE short line:
- [ ] Diagnostic: <1 key suggestion if needed> (e.g., “Focused history for side-specific pain, targeted ROM/strength testing, and basic balance assessment.”)
- [ ] Conservative: <1–2 key rehab strategies> (e.g., “Right-limb strengthening and gait training to reduce asymmetry and double support.”)
- [ ] Interventional: <only if clearly indicated or if red flags are likely> (e.g., “Consider orthopedic referral if unilateral OA or significant structural pathology is suspected and conservative therapy fails.”)
- [ ] Follow-up Metrics: <how to monitor change> (e.g., “Re-assess gait speed, stance/swing symmetry, and double support after 6–8 weeks of therapy.”)

If the available data are limited (e.g., only plantar pressure, only spatiotemporal, or only one test), still follow this structure but keep all sections proportionally brief and avoid speculation beyond the data.
"""


LEGAL_WARNING_TEXT = """
Before proceeding, you MUST confirm the following:

1.  All files you are about to upload have been fully anonymized.
2.  You have removed all Protected Health Information (PHI), including names, dates, MRNs, addresses, and any other identifying data.
3.  You understand that this data will be sent to a third-party AI (OpenAI).
4.  You (or your institution) have the legal and ethical right to upload this anonymized data for analysis.

You are solely responsible for ensuring data privacy compliance.
"""

API_HELP_TEXT = """
How to get an OpenAI API Key

1.  Visit: https://platform.openai.com
2.  Sign in or create a new account.
3.  Add a payment method (Pay-As-You-Go).
    (Note: A ChatGPT Plus subscription is separate and does not include API credits.)
4.  Go to the "API keys" section in your account.
5.  Click "Create new secret key". You’ll get a key that looks like: sk-...
6.  Copy this key and paste it into the previous window.

Keep this key private and do not share it.

If you have trouble, please email:
Mehrdad.Davoudi@med.uni-heidelberg.de
"""

# ----------------------------------------------------------------------
# HELPER FUNCTIONS
# ----------------------------------------------------------------------

def extract_pdf_text_fitz(pdf_path: str) -> str:
    try:
        with fitz.open(pdf_path) as doc:
            text = ""
            for page in doc:
                text += page.get_text()
            return text
    except Exception as e:
        print(f"Error extracting PDF text with fitz: {e}")
        return ""

def extract_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

def extract_docx_images(docx_path: Path) -> list[tuple[bytes, str]]:
    imgs: list[tuple[bytes, str]] = []
    with zipfile.ZipFile(str(docx_path), "r") as z:
        for name in z.namelist():
            if not name.startswith("word/media/"):
                continue
            raw = z.read(name)
            ext = Path(name).suffix
            mime, _ = mimetypes.guess_type(f"x{ext}")
            if not mime:
                mime = "image/png"
            imgs.append((raw, mime))
    return imgs

def to_data_url(mime: str, raw: bytes) -> str:
    return f"data:{mime};base64,{base64.b64encode(raw).decode()}"

def extract_file_contents(path: Path) -> tuple[str, list[tuple[bytes, str]]]:
    text = ""
    images: list[tuple[bytes, str]] = []
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            text = extract_docx_text(path)
            images = extract_docx_images(path)
        elif suffix == ".pdf":
            text = extract_pdf_text_fitz(str(path))
        elif suffix in {".png", ".jpg", ".jpeg", ".bmp"}:
            raw = path.read_bytes()
            mime, _ = mimetypes.guess_type(str(path))
            if not mime:
                mime = "image/png"
            images.append((raw, mime))
        else:
            text = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        print(f"Warning: failed to read {path.name}: {e}")
    return text.strip(), images

def build_content_payload(inputs: list[tuple[str, str, list[tuple[bytes, str]]]]) -> list[dict]:
    content: list[dict] = [
        {"type": "text", "text": CLINICAL_PROMPT.strip()}
    ]
    for text, label, _ in inputs:
        if text:
            content.append({"type": "text", "text": f"FILE [{label}] CONTENTS:\n{text}"})
        else:
            content.append({"type": "text", "text": f"FILE [{label}] had no extractable text."})
    for _, _, imgs in inputs:
        for raw, mime in imgs:
            content.append({"type": "image_url", "image_url": {"url": to_data_url(mime, raw)}})
    return content

def analyze_with_gpt(client: OpenAI, content: list[dict]) -> str:
    response = client.chat.completions.create(
        model="gpt-5.1",
        messages=[{"role": "user", "content": content}],
    )
    return response.choices[0].message.content.strip()

def save_report(header_data: dict, interpretation: str, save_path: str) -> str:
    doc = Document()
    doc.add_heading("Patient Report", level=0)
    doc.add_paragraph(f"Patient Name:\t{header_data['first_name']} {header_data['last_name']}")
    doc.add_paragraph(f"Patient ID:\t{header_data['patient_id']}")
    doc.add_paragraph(f"Date of Birth:\t{header_data['birthdate']}")
    doc.add_heading("Interpretation", level=1)
    for line in interpretation.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    file_name = f"Report_{header_data['last_name']}_{header_data['patient_id']}.docx"
    final_path = os.path.join(save_path, file_name)
    doc.save(final_path)
    return final_path

# ----------------------------------------------------------------------
# ASYNC/THREADED CORE LOGIC
# ----------------------------------------------------------------------

def upload_files_and_create_report(app, files_to_upload, header_data, save_path, api_key) -> None:
    try:
        app.safe_update_status("Connecting to OpenAI...")
        app.safe_update_progress(10)
        client = OpenAI(api_key=api_key)
        client.models.list()
        app.safe_update_status("OpenAI connection successful.")

        app.safe_update_status("Extracting file contents...")
        app.safe_update_progress(25)
        inputs: list[tuple[str, str, list[tuple[bytes, str]]]] = []
        for item in files_to_upload:
            file_path = Path(item["file_path"])
            description = (item.get("description") or "").strip()
            text, images = extract_file_contents(file_path)
            combined_text = f"Description: {description}".strip()
            if text:
                combined_text += f"\n\n{text}"
            inputs.append((combined_text, file_path.name, images))

        app.safe_update_status("Sending data to AI for interpretation (PHI excluded)…")
        app.safe_update_progress(50)
        content_payload = build_content_payload(inputs)
        interpretation = analyze_with_gpt(client, content_payload)
        app.safe_update_status("Interpretation received.")
        app.safe_update_progress(80)

        app.safe_update_status("Saving local report...")
        report_path = save_report(header_data, interpretation, save_path)
        app.safe_update_progress(100)
        app.safe_update_status("Report saved successfully.")

        app.safe_show_info(
            "Process Completed",
            f"Report successfully saved to:\n{report_path}"
        )

    except Exception as e:
        app.safe_update_status(f"Error: {e}")
        app.safe_show_error("Processing Error", f"An error occurred: {e}")
    finally:
        app.safe_update_status("Ready.")
        app.safe_update_progress(0)
        app.toggle_ui_state(True)

# ----------------------------------------------------------------------
# API KEY DIALOG
# ----------------------------------------------------------------------

class APIKeyDialog:
    def __init__(self, parent):
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Enter API Key & Confirm")
        self.dialog.geometry("500x450")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        self.api_key = None

        main_frame = ttk.Frame(self.dialog, padding=15)
        main_frame.pack(fill=tk.BOTH, expand=True)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)

        warning_label = ttk.Label(main_frame, text="Legal & Privacy Warning", font=("Segoe UI", 12, "bold"))
        warning_label.grid(row=0, column=0, sticky=tk.W, pady=(0, 10))

        text_frame = ttk.Frame(main_frame)
        text_frame.grid(row=1, column=0, sticky="nsew", pady=5)
        text_frame.rowconfigure(0, weight=1)
        text_frame.columnconfigure(0, weight=1)
        
        warning_text = scrolledtext.ScrolledText(text_frame, wrap=tk.WORD, height=10, font=("Segoe UI", 11))
        warning_text.grid(row=0, column=0, sticky="nsew")
        warning_text.insert(tk.END, LEGAL_WARNING_TEXT)
        warning_text.config(state=tk.DISABLED)

        key_frame = ttk.Frame(main_frame)
        key_frame.grid(row=2, column=0, sticky="ew", pady=(15, 5))
        key_frame.columnconfigure(1, weight=1)

        ttk.Label(key_frame, text="OpenAI API Key:").grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        self.key_entry = ttk.Entry(key_frame, show="*")
        self.key_entry.grid(row=0, column=1, sticky="ew")

        self.help_btn = ttk.Button(main_frame, text="How do I get an API Key?", command=self.show_api_help, style="Accent.TButton")
        self.help_btn.grid(row=3, column=0, sticky="w", pady=(0, 10))

        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, sticky=tk.E, pady=(10, 0))
        
        self.ok_button = ttk.Button(button_frame, text="OK & Process", command=self.on_ok, style="Accent.TButton")
        self.ok_button.pack(side=tk.LEFT, padx=10)
        self.cancel_button = ttk.Button(button_frame, text="Cancel", command=self.on_cancel)
        self.cancel_button.pack(side=tk.LEFT)

        self.key_entry.focus()
        self.dialog.bind("<Return>", (lambda e: self.on_ok()))

    def show_api_help(self):
        help_win = tk.Toplevel(self.dialog)
        help_win.title("How to get an API Key")
        help_win.geometry("450x300")
        help_win.transient(self.dialog)
        help_win.grab_set()
        
        main = ttk.Frame(help_win, padding=15)
        main.pack(fill=tk.BOTH, expand=True)
        main.rowconfigure(0, weight=1)
        main.columnconfigure(0, weight=1)

        text = scrolledtext.ScrolledText(main, wrap=tk.WORD, font=("Segoe UI", 11))
        text.grid(row=0, column=0, sticky="nsew")
        text.insert(tk.END, API_HELP_TEXT)
        text.config(state=tk.DISABLED)
        
        btn_frame = ttk.Frame(main)
        btn_frame.grid(row=1, column=0, sticky=tk.E, pady=(10, 0))
        ttk.Button(btn_frame, text="Close", command=help_win.destroy).pack()

    def on_ok(self, event=None):
        key = self.key_entry.get().strip()
        if not key:
            messagebox.showwarning("Missing Key", "Please enter your OpenAI API key to proceed.", parent=self.dialog)
            return
        if not (key.startswith("sk-") and len(key) > 50):
            if not messagebox.askyesno("Invalid Key?", "The key format looks incorrect. Proceed anyway?", parent=self.dialog):
                return
                
        self.api_key = key
        self.dialog.destroy()

    def on_cancel(self):
        self.dialog.destroy()

    def show(self):
        self.dialog.wait_window()
        return self.api_key

# ----------------------------------------------------------------------
# DESCRIPTION DIALOG
# ----------------------------------------------------------------------

class DescriptionDialog:
    def __init__(self, parent, filename):
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Add File Description")
        self.dialog.geometry("500x350")
        self.dialog.minsize(400, 300)
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        self.description = None

        main_frame = ttk.Frame(self.dialog, padding=15)
        main_frame.pack(fill=tk.BOTH, expand=True)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)

        label = ttk.Label(main_frame, text=f"Enter a description for:\n{filename}", font=("Segoe UI", 10))
        label.grid(row=0, column=0, sticky=tk.W, pady=(0, 10))

        text_frame = ttk.Frame(main_frame)
        text_frame.grid(row=1, column=0, sticky="nsew", pady=5)
        text_frame.rowconfigure(0, weight=1)
        text_frame.columnconfigure(0, weight=1)
        
        # --- FIXED: Font size increased to 11 ---
        self.text_widget = scrolledtext.ScrolledText(text_frame, wrap=tk.WORD, height=5, font=("Segoe UI", 11))
        self.text_widget.grid(row=0, column=0, sticky="nsew")

        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=2, column=0, sticky=tk.E, pady=(10, 0))
        
        self.ok_button = ttk.Button(button_frame, text="OK", command=self.on_ok, style="Accent.TButton")
        self.ok_button.pack(side=tk.LEFT, padx=10)
        self.cancel_button = ttk.Button(button_frame, text="Cancel", command=self.on_cancel)
        self.cancel_button.pack(side=tk.LEFT)

        self.text_widget.focus()
        self.dialog.bind("<Return>", (lambda e: self.on_ok()))

    def on_ok(self, event=None):
        self.description = self.text_widget.get("1.0", tk.END).strip()
        self.dialog.destroy()

    def on_cancel(self):
        self.description = None
        self.dialog.destroy()

    def show(self):
        self.dialog.wait_window()
        return self.description

# ----------------------------------------------------------------------
# MAIN APP GUI
# ----------------------------------------------------------------------

class FileCollectorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AI Clinical Interpreter")
        self.root.geometry("850x820")
        self.root.minsize(800, 750)

        self.title_font = ("Segoe UI", 18, "bold")
        self.subtitle_font = ("Segoe UI", 10, "bold")
        self.label_font = ("Segoe UI", 10, "bold")
        self.base_font = ("Segoe UI", 10, "bold") 

        self.file_list: list[dict] = []
        self.save_location = ""
        self.interactive_widgets = []

        self.help_window = HelpWindow(root)
        self.ack_window = AcknowledgmentWindow(root)
        self.license_window = LicenseWindow(root)

        self.load_user_photo()
        self.build_menubar()
        self.create_widgets()

    def load_user_photo(self):
        img = None
        for p in USER_PHOTO_FILES:
            if os.path.exists(p):
                try:
                    img = Image.open(p)
                    break
                except Exception:
                    continue
        if img is None:
            img = Image.new("RGB", (100, 100), color="#cccccc")

        img_large = img.resize((90, 90), Image.LANCZOS)
        self.user_photo_large = ImageTk.PhotoImage(img_large)
        self.root.user_photo_large = self.user_photo_large
        
        img_small = img.resize((50, 50), Image.LANCZOS)
        self.user_photo_small = ImageTk.PhotoImage(img_small)
        self.root.user_photo_small = self.user_photo_small
        
        try:
            self.root.iconphoto(True, self.user_photo_large)
            if os.path.exists(USER_ICON_FILE):
                self.root.iconbitmap(USER_ICON_FILE)
        except Exception as e:
            print(f"Warning: Could not set icon. {e}")

    def build_menubar(self):
        menubar = tk.Menu(self.root)

        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="Usage Help", command=self.help_window.show)
        menubar.add_cascade(label="Help", menu=help_menu)

        about_menu = tk.Menu(menubar, tearoff=0)
        about_menu.add_command(
            label="View License", command=self.license_window.show
        )
        about_menu.add_command(
            label="Acknowledgments", command=self.ack_window.show
        )
        menubar.add_cascade(label="About", menu=about_menu)

        self.root.config(menu=menubar)

    def create_widgets(self):
        main = ttk.Frame(self.root, padding="20")
        main.pack(fill=tk.BOTH, expand=True)
        main.columnconfigure(0, weight=1)
        main.rowconfigure(3, weight=1) 

        # HEADER
        header = ttk.Frame(main)
        header.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 25))
        header.columnconfigure(1, weight=1)

        self.photo_label = ttk.Label(header, image=self.user_photo_large)
        self.photo_label.grid(row=0, column=0, rowspan=4, sticky=tk.NW, padx=(0, 20))
        self.photo_label.image = self.user_photo_large

        title_label = ttk.Label(
            header, text="AI Clinical Interpreter", font=self.title_font
        )
        title_label.grid(row=0, column=1, sticky=tk.W)

        desc1 = ttk.Label(
            header,
            text="Mehrdad Davoudi",
            font=self.subtitle_font,
        )
        desc1.grid(row=1, column=1, sticky=tk.W, pady=(5,0))

        desc2 = ttk.Label(
            header,
            text=(
                "PhD student, Clinic for Orthopaedics, "
                "Heidelberg University Hospital, Heidelberg, Germany."
            ),
            font=self.base_font
        )
        desc2.grid(row=2, column=1, sticky=tk.W)

        desc3 = ttk.Label(
            header,
            text="Email: Mehrdad.Davoudi@med.uni-heidelberg.de",
            font=self.base_font
        )
        desc3.grid(row=3, column=1, sticky=tk.W, pady=(2, 0))


        # Patient info
        self.patient_frame = ttk.LabelFrame(main, text="1. Patient Information")  
        self.patient_frame.grid(row=1, column=0, pady=10, sticky="ew")
        self.patient_frame.grid_columnconfigure(1, weight=1)

        ttk.Label(self.patient_frame, text="First Name:", font=self.label_font).grid(row=0, column=0, padx=10, pady=5, sticky="w")
        self.first_name_entry = ttk.Entry(self.patient_frame); self.first_name_entry.grid(row=0, column=1, padx=10, pady=5, sticky="ew")
        ttk.Label(self.patient_frame, text="Last Name:", font=self.label_font).grid(row=1, column=0, padx=10, pady=5, sticky="w")
        self.last_name_entry = ttk.Entry(self.patient_frame); self.last_name_entry.grid(row=1, column=1, padx=10, pady=5, sticky="ew")
        ttk.Label(self.patient_frame, text="Patient ID:", font=self.label_font).grid(row=2, column=0, padx=10, pady=5, sticky="w")
        self.id_entry = ttk.Entry(self.patient_frame); self.id_entry.grid(row=2, column=1, padx=10, pady=5, sticky="ew")
        ttk.Label(self.patient_frame, text="Birthdate:", font=self.label_font).grid(row=3, column=0, padx=10, pady=5, sticky="w")
        self.birthdate_entry = ttk.Entry(self.patient_frame); self.birthdate_entry.grid(row=3, column=1, padx=10, pady=5, sticky="ew")

        # Save location
        self.save_frame = ttk.LabelFrame(main, text="2. Save Location")
        self.save_frame.grid(row=2, column=0, pady=10, sticky="ew")
        self.save_frame.columnconfigure(0, weight=1)

        self.save_btn = ttk.Button(self.save_frame, text="Select Save Location", command=self.select_save_location, style="Accent.TButton")
        self.save_btn.grid(row=0, column=0, padx=10, pady=10, sticky="ew")
        self.save_label = ttk.Label(self.save_frame, text="No location selected", style="Secondary.TLabel", width=30, anchor=tk.W)
        self.save_label.grid(row=0, column=1, padx=10, pady=10, sticky="w")

        # Files
        self.file_frame = ttk.LabelFrame(main, text="3. Add Anonymized Files")
        self.file_frame.grid(row=3, column=0, pady=10, sticky="nsew")
        self.file_frame.rowconfigure(1, weight=1)
        self.file_frame.columnconfigure(0, weight=1)

        self.button_frame = ttk.Frame(self.file_frame); self.button_frame.grid(row=0, column=0, pady=10, padx=10, sticky="w")
        self.add_file_btn = ttk.Button(self.button_frame, text="Add File", command=self.add_file, style="Accent.TButton")
        self.add_file_btn.pack(side=tk.LEFT, padx=5)

        self.file_display = scrolledtext.ScrolledText(self.file_frame, width=74, height=10, font=("Segoe UI", 10))
        self.file_display.grid(row=1, column=0, pady=10, padx=10, sticky="nsew")
        self.file_display.insert(tk.END, "Click 'Add File' to begin."); self.file_display.config(state=tk.DISABLED)

        # Execute
        self.execute_frame = ttk.LabelFrame(main, text="4. Process Report")
        self.execute_frame.grid(row=4, column=0, pady=10, sticky="ew")
        self.execute_frame.columnconfigure(0, weight=1)
        self.execute_frame.columnconfigure(1, weight=1)
        
        self.finish_btn = ttk.Button(self.execute_frame, text="Finish & Process", command=self.finish_and_process, style="Accent.TButton")
        self.finish_btn.grid(row=0, column=0, padx=10, pady=10, sticky="ew", ipady=5)
        
        # Reset button
        self.reset_btn = ttk.Button(self.execute_frame, text="Clear Form for New Patient", command=self.reset_form)
        self.reset_btn.grid(row=0, column=1, padx=10, pady=10, sticky="ew", ipady=5)

        # Status Bar
        status = ttk.Frame(main, padding="5 0 0 0")
        status.grid(row=5, column=0, sticky="ew")
        status.columnconfigure(0, weight=1)
        self.status_text = tk.StringVar(value="Ready.")
        self.status_label = ttk.Label(status, textvariable=self.status_text, anchor=tk.W)
        self.status_label.grid(row=0, column=0, sticky="ew")
        self.progress = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(status, variable=self.progress, mode="determinate")
        self.progress_bar.grid(row=1, column=0, sticky="ew", pady=(5,0))

        self.interactive_widgets = [
            self.first_name_entry, self.last_name_entry, self.id_entry,
            self.birthdate_entry, self.save_btn, self.add_file_btn, 
            self.finish_btn, self.reset_btn
        ]

    def safe_update_status(self, text):
        self.root.after(0, lambda: self.status_text.set(text))

    def safe_update_progress(self, value):
        self.root.after(0, lambda: self.progress.set(value))

    def safe_show_error(self, title, message):
        self.root.after(0, lambda: messagebox.showerror(title, message))
        
    def safe_show_info(self, title, message):
        self.root.after(0, lambda: messagebox.showinfo(title, message))

    def toggle_ui_state(self, enabled):
        state = tk.NORMAL if enabled else tk.DISABLED
        for w in self.interactive_widgets:
            try:
                w.config(state=state)
            except tk.TclError:
                pass

    def reset_form(self):
        self.first_name_entry.delete(0, tk.END)
        self.last_name_entry.delete(0, tk.END)
        self.id_entry.delete(0, tk.END)
        self.birthdate_entry.delete(0, tk.END)
        
        self.save_location = ""
        self.save_label.config(text="No location selected", style="Secondary.TLabel")
        
        self.file_list = []
        self.update_display()
        
        self.status_text.set("Ready.")
        self.progress.set(0)
        self.toggle_ui_state(True)

    def select_save_location(self) -> None:
        path = filedialog.askdirectory(title="Select folder to save report")
        if path:
            self.save_location = path
            self.save_label.config(text=f"Saving to: …{os.path.basename(path)}", style="TLabel")

    def add_file(self) -> None:
        file_path = filedialog.askopenfilename(title="Select a file (PDF/DOCX/PNG/JPG/BMP/TXT)")
        if not file_path:
            return
        
        filename = os.path.basename(file_path)
        dialog = DescriptionDialog(self.root, filename)
        description = dialog.show()

        if description is None:
            return
            
        self.file_list.append({"file_path": file_path, "description": description})
        self.update_display()

    def update_display(self) -> None:
        self.file_display.config(state=tk.NORMAL); self.file_display.delete("1.0", tk.END)
        if not self.file_list:
            self.file_display.insert(tk.END, "Click 'Add File' to begin.")
        else:
            for i, item in enumerate(self.file_list, 1):
                self.file_display.insert(
                    tk.END,
                    f"File {i}:\n  Path: {item['file_path']}\n  Desc: {item['description']}\n\n"
                )
        self.file_display.config(state=tk.DISABLED)

    def finish_and_process(self) -> None:
        header_data = {
            "first_name": self.first_name_entry.get().strip(),
            "last_name": self.last_name_entry.get().strip(),
            "patient_id": self.id_entry.get().strip(),
            "birthdate": self.birthdate_entry.get().strip(),
        }
        if not all(header_data.values()):
            messagebox.showwarning("Missing Info", "Please fill out all patient header fields.")
            return
        if not self.save_location:
            messagebox.showwarning("Missing Info", "Please select a save location for the report.")
            return
        if not self.file_list:
            messagebox.showwarning("No Files", "No files were added. Please add files first.")
            return
            
        dialog = APIKeyDialog(self.root)
        api_key = dialog.show()
        
        if not api_key:
            self.safe_update_status("Processing cancelled (no API key).")
            return

        print("--- GUI finished. Starting backend processing (PHI excluded from model)… ---")
        
        self.toggle_ui_state(False)
        self.safe_update_status("Starting processing...")
        th = threading.Thread(
            target=upload_files_and_create_report,
            args=(self, self.file_list, header_data, self.save_location, api_key),
            daemon=True
        )
        th.start()
        
# ----------------------------------------------------------------------
# MAIN
# ----------------------------------------------------------------------

def main() -> None:
    print("Launching AI Clinical Interpreter GUI…")
    root = tk.Tk()
    sv_ttk.set_theme("light")
    style = ttk.Style()
    style.configure("TLabelframe.Label", font=("Segoe UI", 11))
    app = FileCollectorApp(root)
    root.mainloop()
    print("GUI closed. Program finished.")

if __name__ == "__main__":
    main()